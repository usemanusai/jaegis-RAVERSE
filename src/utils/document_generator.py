"""
Document generation utilities for RAVERSE 2.0.
Supports DOCX, PDF, and Markdown output formats.
Replaces proprietary Microsoft Word with open-source alternatives.
"""

import logging
import os
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """
    Generate research reports in multiple formats (DOCX, PDF, Markdown).
    Uses python-docx for DOCX generation and Pandoc for format conversion.
    """

    def __init__(self, output_dir: str = "./reports"):
        """
        Initialize document generator.

        Args:
            output_dir: Directory for output files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logging.getLogger(f"RAVERSE.DocumentGenerator")

    def generate_docx(
        self,
        title: str,
        content: Dict[str, Any],
        filename: Optional[str] = None,
    ) -> str:
        """
        Generate DOCX document from research data.

        Args:
            title: Document title
            content: Dictionary with sections and content
            filename: Output filename (auto-generated if None)

        Returns:
            Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        filename = filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename

        try:
            doc = Document()

            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
            meta_para.add_run(f"Format: DOCX (Open Document)\n").italic = True

            # Add sections
            for section_name, section_data in content.items():
                if isinstance(section_data, dict):
                    doc.add_heading(section_name, level=1)
                    for key, value in section_data.items():
                        doc.add_heading(key, level=2)
                        if isinstance(value, list):
                            for item in value:
                                doc.add_paragraph(str(item), style="List Bullet")
                        else:
                            doc.add_paragraph(str(value))
                else:
                    doc.add_heading(section_name, level=1)
                    doc.add_paragraph(str(section_data))

            # Save document
            doc.save(filepath)
            self.logger.info(f"DOCX report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            self.logger.error(f"Failed to generate DOCX: {e}")
            raise

    def generate_markdown(
        self,
        title: str,
        content: Dict[str, Any],
        filename: Optional[str] = None,
    ) -> str:
        """
        Generate Markdown document from research data.

        Args:
            title: Document title
            content: Dictionary with sections and content
            filename: Output filename (auto-generated if None)

        Returns:
            Path to generated Markdown file
        """
        filename = filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        filepath = self.output_dir / filename

        try:
            lines = [
                f"# {title}\n",
                f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n",
                f"**Format:** Markdown\n\n",
            ]

            for section_name, section_data in content.items():
                lines.append(f"## {section_name}\n")
                if isinstance(section_data, dict):
                    for key, value in section_data.items():
                        lines.append(f"### {key}\n")
                        if isinstance(value, list):
                            for item in value:
                                lines.append(f"- {item}\n")
                        else:
                            lines.append(f"{value}\n")
                else:
                    lines.append(f"{section_data}\n")
                lines.append("\n")

            with open(filepath, "w", encoding="utf-8") as f:
                f.writelines(lines)

            self.logger.info(f"Markdown report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            self.logger.error(f"Failed to generate Markdown: {e}")
            raise

    def generate_pdf(
        self,
        title: str,
        content: Dict[str, Any],
        filename: Optional[str] = None,
    ) -> str:
        """
        Generate PDF document using Pandoc (requires Markdown first).

        Args:
            title: Document title
            content: Dictionary with sections and content
            filename: Output filename (auto-generated if None)

        Returns:
            Path to generated PDF file
        """
        if not PANDOC_AVAILABLE:
            raise ImportError("pypandoc not installed. Install with: pip install pypandoc")

        # First generate Markdown
        md_filename = f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        md_path = self.generate_markdown(title, content, md_filename)

        filename = filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = self.output_dir / filename

        try:
            # Convert Markdown to PDF using Pandoc
            pypandoc.convert_file(
                md_path,
                "pdf",
                outputfile=str(filepath),
                extra_args=["--pdf-engine=xelatex"],
            )

            # Clean up temporary Markdown file
            Path(md_path).unlink()

            self.logger.info(f"PDF report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            self.logger.error(f"Failed to generate PDF: {e}")
            raise

    def generate_all_formats(
        self,
        title: str,
        content: Dict[str, Any],
        base_filename: Optional[str] = None,
    ) -> Dict[str, str]:
        """
        Generate reports in all available formats.

        Args:
            title: Document title
            content: Dictionary with sections and content
            base_filename: Base filename (without extension)

        Returns:
            Dictionary mapping format to filepath
        """
        base_filename = base_filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        results = {}

        # Generate DOCX
        if DOCX_AVAILABLE:
            try:
                results["docx"] = self.generate_docx(title, content, f"{base_filename}.docx")
            except Exception as e:
                self.logger.warning(f"DOCX generation failed: {e}")

        # Generate Markdown
        try:
            results["markdown"] = self.generate_markdown(title, content, f"{base_filename}.md")
        except Exception as e:
            self.logger.warning(f"Markdown generation failed: {e}")

        # Generate PDF
        if PANDOC_AVAILABLE:
            try:
                results["pdf"] = self.generate_pdf(title, content, f"{base_filename}.pdf")
            except Exception as e:
                self.logger.warning(f"PDF generation failed: {e}")

        return results


# Convenience functions
def generate_research_report(
    title: str,
    research_data: Dict[str, Any],
    output_format: str = "docx",
    output_dir: str = "./reports",
) -> str:
    """
    Generate a research report in specified format.

    Args:
        title: Report title
        research_data: Research content
        output_format: Format (docx, markdown, pdf)
        output_dir: Output directory

    Returns:
        Path to generated report
    """
    generator = DocumentGenerator(output_dir)

    if output_format.lower() == "docx":
        return generator.generate_docx(title, research_data)
    elif output_format.lower() == "markdown":
        return generator.generate_markdown(title, research_data)
    elif output_format.lower() == "pdf":
        return generator.generate_pdf(title, research_data)
    else:
        raise ValueError(f"Unsupported format: {output_format}")

